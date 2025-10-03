"""Document parser for extracting text from various file formats."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents and extract text content."""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")

            content = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(reader.pages)} pages from PDF: {file_path}")
            return content

        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise ValueError(f"Failed to parse PDF: {e}") from e

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from Word document (.docx).
        Works with both Microsoft Word and Google Docs exports.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = [f"--- Table {table_num} ---"]
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                text_parts.append("\n".join(table_text))

            content = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables from DOCX: {file_path}")
            return content

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise ValueError(f"Failed to parse Word document: {e}") from e

    @staticmethod
    def parse_xlsx(file_path: str) -> str:
        """
        Extract text from Excel spreadsheet (.xlsx).
        Works with both Microsoft Excel and Google Sheets exports.

        Args:
            file_path: Path to XLSX file

        Returns:
            Extracted text content with context
        """
        try:
            workbook = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"=== Sheet: {sheet_name} ==="]

                # Get header row (first row)
                headers = []
                first_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
                if first_row:
                    headers = [str(h) if h is not None else f"Column{i+1}"
                              for i, h in enumerate(first_row)]

                # Process data rows
                row_count = 0
                for row_num, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), 2):
                    # Skip empty rows
                    if not any(cell is not None and str(cell).strip() for cell in row):
                        continue

                    # Create contextual text for each row
                    row_text = []
                    for col_idx, cell_value in enumerate(row):
                        if cell_value is not None and str(cell_value).strip():
                            header = headers[col_idx] if col_idx < len(headers) else f"Column{col_idx+1}"
                            row_text.append(f"{header}: {cell_value}")

                    if row_text:
                        sheet_text.append(f"Row {row_num}: " + ", ".join(row_text))
                        row_count += 1

                if row_count > 0:
                    text_parts.append("\n".join(sheet_text))

            content = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(workbook.sheetnames)} sheets from XLSX: {file_path}")
            return content

        except Exception as e:
            logger.error(f"Failed to parse XLSX {file_path}: {e}")
            raise ValueError(f"Failed to parse Excel file: {e}") from e

    @staticmethod
    def parse_text(file_path: str) -> str:
        """
        Read plain text file.

        Args:
            file_path: Path to text file

        Returns:
            File content
        """
        try:
            with open(file_path, encoding='utf-8') as f:
                content = f.read()
            logger.info(f"Read text file: {file_path}")
            return content

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, encoding='latin-1') as f:
                    content = f.read()
                logger.info(f"Read text file with latin-1 encoding: {file_path}")
                return content
            except Exception as e:
                logger.error(f"Failed to parse text file {file_path}: {e}")
                raise ValueError(f"Failed to parse text file: {e}") from e
        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            raise ValueError(f"Failed to parse text file: {e}") from e

    @classmethod
    def parse_document(cls, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Parse document and extract text based on file type.

        Args:
            file_path: Path to document
            file_type: Optional file extension (auto-detected if not provided)

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported or parsing fails
        """
        path = Path(file_path)

        if not path.exists():
            raise ValueError(f"File not found: {file_path}")

        # Determine file type
        if file_type is None:
            file_type = path.suffix.lower()
        else:
            file_type = file_type.lower()
            if not file_type.startswith('.'):
                file_type = f'.{file_type}'

        # Parse based on file type
        if file_type == '.pdf':
            return cls.parse_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            if file_type == '.doc':
                raise ValueError("Legacy .doc format not supported. Please save as .docx")
            return cls.parse_docx(file_path)
        elif file_type in ['.xlsx', '.xls']:
            if file_type == '.xls':
                raise ValueError("Legacy .xls format not supported. Please save as .xlsx")
            return cls.parse_xlsx(file_path)
        elif file_type in ['.txt', '.md', '.csv', '.json', '.log']:
            return cls.parse_text(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported formats: PDF (.pdf), Word (.docx), Excel (.xlsx), "
                f"Text (.txt, .md, .csv, .json, .log)"
            )
